from docx import Document

def convert_docx_to_md(docx_file, md_file):
    """Converts a given .docx file to a .md file, attempting to preserve basic formatting."""

    doc = Document(docx_file)

    with open(md_file, 'w', encoding='utf-8') as f:
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                f.write(f"{text}\n\n")

            # Handle bold and italic formatting
            for run in paragraph.runs:
                if run.bold:
                    f.write("**")
                if run.italic:
                    f.write("*")
                f.write(run.text)
                if run.bold:
                    f.write("**")
                if run.italic:
                    f.write("*")
            f.write("\n\n")

        # Handle tables (basic conversion)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    f.write(cell.text + " | ")
                f.write("\n")
            f.write("\n")

# Example usage:
docx_file = "21101-840.docx"
md_file = "output.md"
convert_docx_to_md(docx_file, md_file)